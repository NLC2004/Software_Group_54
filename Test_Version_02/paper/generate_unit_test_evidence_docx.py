from pathlib import Path
import re

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MARKDOWN = ROOT / "UNIT_TEST_EVIDENCE.md"
ASSETS = ROOT / "document_assets"
OUTPUT = ROOT / "Unit_Test_Evidence_Professional.docx"
SUMMARY = ASSETS / "unit_test_summary.png"
FLOW = ASSETS / "unit_test_method_flow.png"

NAVY = "122F47"
BLUE = "2563EB"
TEAL = "0F766E"
INK = "1F2937"
MUTED = "64748B"
PALE = "EFF6FF"
PALE_ALT = "F8FAFC"
BORDER = "D9E3F0"


def font(size, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def rounded(draw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def build_summary_image():
    image = Image.new("RGB", (1800, 900), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(52, True)
    label_font = font(22, True)
    metric_font = font(72, True)
    small_font = font(20)
    bar_font = font(25, True)

    draw.rectangle((0, 0, 1800, 16), fill="#" + BLUE)
    draw.text((90, 65), "AUTOMATED TEST RESULT", fill="#" + NAVY, font=title_font)
    rounded(draw, (1405, 63, 1710, 132), 34, "#" + TEAL)
    draw.text((1503, 86), "PASS", fill="white", font=label_font)

    metrics = [
        ("103", "TESTS PASSED", BLUE),
        ("0", "TESTS FAILED", TEAL),
        ("18", "TEST CLASSES", NAVY),
    ]
    for idx, (value, label, colour) in enumerate(metrics):
        left = 90 + idx * 550
        rounded(draw, (left, 192, left + 500, 407), 20, "#F8FAFC", "#" + BORDER, 2)
        draw.text((left + 38, 230), value, fill="#" + colour, font=metric_font)
        draw.text((left + 40, 333), label, fill="#" + MUTED, font=label_font)

    draw.text((90, 494), "Coverage by test ownership area", fill="#" + NAVY, font=label_font)
    areas = [("TA workflows", 38, BLUE), ("Administrator", 28, NAVY), ("MO workflows", 19, TEAL),
             ("Shared / service / AI", 18, "475569")]
    max_value = 38
    for idx, (name, value, colour) in enumerate(areas):
        y = 570 + idx * 70
        draw.text((90, y), name, fill="#" + INK, font=small_font)
        rounded(draw, (430, y + 2, 1460, y + 36), 17, "#E8EEF6")
        rounded(draw, (430, y + 2, 430 + int(1030 * value / max_value), y + 36), 17, "#" + colour)
        draw.text((1505, y + 5), str(value), fill="#" + NAVY, font=bar_font)

    image.save(SUMMARY)


def build_flow_image():
    image = Image.new("RGB", (1800, 560), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(42, True)
    box_font = font(25, True)
    note_font = font(18)

    draw.text((75, 48), "HOW THE AUTOMATED TESTS VERIFY THE SYSTEM", fill="#" + NAVY, font=title_font)
    boxes = [
        ("JUnit 5 Case", "Arrange inputs\nand test identities"),
        ("TestHttpExchange", "Simulate API\nHTTP request"),
        ("Handlers + Services", "Run real rules\nand permissions"),
        ("Temp JSON Evidence", "Assert records,\nfiles and response"),
    ]
    for idx, (heading, note) in enumerate(boxes):
        x = 70 + idx * 430
        rounded(draw, (x, 178, x + 345, 394), 18, "#F8FAFC", "#" + BORDER, 2)
        draw.rectangle((x, 178, x + 345, 188), fill="#" + BLUE)
        draw.text((x + 28, 224), heading, fill="#" + NAVY, font=box_font)
        draw.multiline_text((x + 28, 280), note, fill="#" + MUTED, font=note_font, spacing=10)
        if idx < len(boxes) - 1:
            ax = x + 365
            draw.line((ax, 286, ax + 42, 286), fill="#" + BLUE, width=5)
            draw.polygon([(ax + 42, 274), (ax + 62, 286), (ax + 42, 298)], fill="#" + BLUE)
    image.save(FLOW)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "5")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margin(cell, top=90, start=115, bottom=90, end=115):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn("w:" + m))
        if node is None:
            node = OxmlElement("w:" + m)
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_run_font(run, name="Aptos", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_field(paragraph, instruction, placeholder=""):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def add_inline(paragraph, text, size=None, color=None):
    pieces = re.split(r"(`[^`]+`)", text)
    for piece in pieces:
        if not piece:
            continue
        if piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            set_run_font(run, "Consolas", size or 9, BLUE)
        else:
            run = paragraph.add_run(piece)
            set_run_font(run, "Aptos", size, color)


def style_document(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.14
    for style_name, size, colour in (("Heading 1", 20, NAVY), ("Heading 2", 14, BLUE), ("Heading 3", 11, NAVY)):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(colour)
        style.paragraph_format.space_before = Pt(15 if style_name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True


def setup_section(section, landscape=False, first=False):
    section.orientation = WD_ORIENT.LANDSCAPE if landscape else WD_ORIENT.PORTRAIT
    if landscape:
        section.page_width, section.page_height = Cm(29.7), Cm(21.0)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)
    else:
        section.page_width, section.page_height = Cm(21.0), Cm(29.7)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.75)
    if first:
        section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.clear()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("SOFTWARE GROUP 54   /   UNIT TEST EVIDENCE")
    set_run_font(run, size=8, color=MUTED, bold=True)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Test_Version_02    |    Verified 24 May 2026    |    Page ")
    set_run_font(run, size=8, color=MUTED)
    add_field(fp, "PAGE", "1")


def add_cover(document):
    band = document.add_table(rows=1, cols=2)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    band.cell(0, 0).width = Cm(12.4)
    band.cell(0, 1).width = Cm(4.6)
    for cell in band.row_cells(0):
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY)
        set_cell_margin(cell, 130, 180, 130, 180)
    band.cell(0, 0).text = ""
    r = band.cell(0, 0).paragraphs[0].add_run("GROUP 54")
    set_run_font(r, size=12, color="FFFFFF", bold=True)
    band.cell(0, 1).text = ""
    p = band.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("TEST EVIDENCE")
    set_run_font(r, size=8, color="D6E4F0", bold=True)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(68)

    marker = document.add_paragraph()
    marker.paragraph_format.space_after = Pt(14)
    r = marker.add_run("SOFTWARE ENGINEERING  /  AUTOMATED VERIFICATION")
    set_run_font(r, size=9, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(0)
    run = title.add_run("Unit Test")
    set_run_font(run, "Aptos Display", 41, NAVY, True)
    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(14)
    run = title.add_run("Evidence Report")
    set_run_font(run, "Aptos Display", 41, BLUE, True)

    rule = document.add_table(rows=1, cols=1)
    rule.alignment = WD_TABLE_ALIGNMENT.LEFT
    rule.autofit = False
    rule.cell(0, 0).width = Cm(3.0)
    set_cell_shading(rule.cell(0, 0), BLUE)
    set_cell_border(rule.cell(0, 0), BLUE)
    set_cell_margin(rule.cell(0, 0), 25, 0, 25, 0)
    rule.cell(0, 0).text = ""

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(22)
    subtitle.paragraph_format.space_after = Pt(5)
    run = subtitle.add_run("BUPT International School")
    set_run_font(run, size=15, color=NAVY, bold=True)
    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(82)
    run = subtitle.add_run("Teaching Assistant Recruitment System")
    set_run_font(run, size=15, color=INK)

    status = document.add_table(rows=1, cols=2)
    status.alignment = WD_TABLE_ALIGNMENT.CENTER
    status.autofit = False
    status.cell(0, 0).width = Cm(11.2)
    status.cell(0, 1).width = Cm(5.8)
    for cell in status.row_cells(0):
        set_cell_border(cell, BORDER)
        set_cell_margin(cell, 165, 170, 165, 170)
    set_cell_shading(status.cell(0, 0), PALE_ALT)
    set_cell_shading(status.cell(0, 1), PALE)
    status.cell(0, 0).text = ""
    p = status.cell(0, 0).paragraphs[0]
    r = p.add_run("REPORT SCOPE\n")
    set_run_font(r, size=7, color=MUTED, bold=True)
    r = p.add_run("JUnit 5 backend and handler-level verification")
    set_run_font(r, size=10, color=INK, bold=True)
    status.cell(0, 1).text = ""
    p = status.cell(0, 1).paragraphs[0]
    r = p.add_run("RESULT\n")
    set_run_font(r, size=7, color=MUTED, bold=True)
    r = p.add_run("103 / 103  PASS")
    set_run_font(r, size=13, color=TEAL, bold=True)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(28)

    meta = document.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    values = [("TESTED VERSION", "Test_Version_02"), ("FRAMEWORK", "JUnit 5"), ("VERIFIED DATE", "24 May 2026")]
    for idx, (label, value) in enumerate(values):
        cell = meta.cell(0, idx)
        cell.width = Cm(5.55)
        set_cell_shading(cell, PALE_ALT)
        set_cell_border(cell, BORDER)
        set_cell_margin(cell, 150, 150, 150, 150)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(label + "\n")
        set_run_font(r, size=7, color=MUTED, bold=True)
        r = p.add_run(value)
        set_run_font(r, size=11, color=NAVY, bold=True)

    note = document.add_paragraph()
    note.paragraph_format.space_before = Pt(34)
    note.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = note.add_run("Prepared by Group 54  |  Academic software engineering assessment")
    set_run_font(run, size=8, color=MUTED)


def add_contents(document):
    document.add_page_break()
    title = document.add_paragraph(style="Heading 1")
    title.add_run("Contents")
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    add_inline(p, "Automatic table of contents. Open in Microsoft Word and update fields if page numbers require refresh.", 9, MUTED)
    toc = document.add_paragraph()
    toc.paragraph_format.space_after = Pt(22)
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u', "Update field in Word to populate the table of contents.")

    document.add_paragraph("Document Map", style="Heading 2")
    items = [
        ("01-04", "Control, summary, strategy and environment"),
        ("05-06", "Execution instructions and user-story coverage"),
        ("07-08", "Additional feature coverage and class inventory"),
        ("09", "Complete 103-case unit test design catalogue"),
        ("10", "Selected critical test design details"),
        ("11-13", "Execution evidence, limitations and security observations"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (number, text) in enumerate(items):
        cells = table.add_row().cells
        cells[0].width = Cm(2.4)
        cells[1].width = Cm(13.6)
        for cell in cells:
            set_cell_border(cell, BORDER)
            set_cell_margin(cell, 115, 140, 115, 140)
            set_cell_shading(cell, "FFFFFF" if index % 2 == 0 else PALE_ALT)
        cells[0].text = number
        cells[1].text = text
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, size=9, color=BLUE, bold=True)
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, size=9, color=INK)


def table_widths(col_count, landscape):
    if landscape and col_count == 4:
        return [0.75, 2.35, 4.00, 3.25]
    if col_count == 4:
        return [0.62, 1.18, 2.50, 2.35]
    if col_count == 3:
        return [2.10, 0.65, 3.85]
    if col_count == 2:
        return [2.00, 4.60]
    return [6.8 / col_count] * col_count


def add_markdown_table(document, rows, landscape):
    if len(rows) < 2:
        return
    columns = len(rows[0])
    table = document.add_table(rows=0, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = table_widths(columns, landscape)
    body_size = 7.0 if landscape and columns == 4 else 8.1
    for index, source_row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        if index == 0:
            repeat_table_header(table.rows[-1])
        for column, value in enumerate(source_row):
            cells[column].width = Inches(widths[column])
            cells[column].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cells[column], 72 if landscape else 90, 88, 72 if landscape else 90, 88)
            set_cell_border(cells[column])
            set_cell_shading(cells[column], NAVY if index == 0 else ("FFFFFF" if index % 2 else PALE_ALT))
            cells[column].text = ""
            p = cells[column].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            add_inline(p, value.strip(), body_size, "FFFFFF" if index == 0 else INK)
            for run in p.runs:
                if index == 0:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code_block(document, text):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, NAVY)
    set_cell_border(cell, NAVY)
    set_cell_margin(cell, 135, 150, 135, 150)
    cell.text = ""
    for line_number, line in enumerate(text.splitlines()):
        p = cell.paragraphs[0] if line_number == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        set_run_font(run, "Consolas", 7.8, "F8FAFC")
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    table_lines = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for index, line in enumerate(table_lines):
        values = [part.strip() for part in line.strip("|").split("|")]
        if index == 1 and all(re.match(r"^:?-{2,}:?$", item) for item in values):
            continue
        rows.append(values)
    return rows, i


def render_markdown(document):
    lines = MARKDOWN.read_text(encoding="utf-8").splitlines()
    i = 0
    started = False
    landscape = False
    pending = []

    def flush_paragraph():
        nonlocal pending
        if pending:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(7)
            add_inline(p, " ".join(item.strip() for item in pending), 9.2, INK)
            pending = []

    while i < len(lines):
        line = lines[i]
        if not started:
            if line.startswith("## 1."):
                started = True
            else:
                i += 1
                continue
        if line.startswith("## 9."):
            flush_paragraph()
            section = document.add_section(WD_SECTION.NEW_PAGE)
            setup_section(section, landscape=True)
            landscape = True
        elif line.startswith("## 10."):
            flush_paragraph()
            section = document.add_section(WD_SECTION.NEW_PAGE)
            setup_section(section, landscape=False)
            landscape = False
        if line.startswith("## "):
            flush_paragraph()
            if line.startswith("## 3."):
                document.add_picture(str(SUMMARY), width=Cm(16.7))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = document.add_paragraph("Figure 1. Verified test execution overview")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(cap.runs[0], size=8, color=MUTED)
            if line.startswith("## 4."):
                document.add_picture(str(FLOW), width=Cm(16.7))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = document.add_paragraph("Figure 2. Handler-level unit testing method")
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(cap.runs[0], size=8, color=MUTED)
            document.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            flush_paragraph()
            document.add_heading(line[4:].strip(), level=2)
            i += 1
            continue
        if line.startswith("```"):
            flush_paragraph()
            code = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            add_code_block(document, "\n".join(code))
            i += 1
            continue
        if line.lstrip().startswith("|"):
            flush_paragraph()
            rows, i = parse_table(lines, i)
            add_markdown_table(document, rows, landscape)
            continue
        if line.startswith("- "):
            flush_paragraph()
            p = document.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_inline(p, line[2:].strip(), 9.1, INK)
            i += 1
            continue
        if not line.strip():
            flush_paragraph()
            i += 1
            continue
        pending.append(line)
        i += 1
    flush_paragraph()


def build_document():
    ASSETS.mkdir(exist_ok=True)
    build_summary_image()
    build_flow_image()
    document = Document()
    style_document(document)
    first = document.sections[0]
    setup_section(first, first=True)
    settings = document.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    document.core_properties.title = "Unit Test Evidence Report - Test_Version_02"
    document.core_properties.subject = "JUnit 5 automated testing evidence"
    document.core_properties.author = "Software Group 54"
    add_cover(document)
    add_contents(document)
    render_markdown(document)
    document.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_document()
